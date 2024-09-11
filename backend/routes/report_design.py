from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_style(cell, font_name, font_size, bold=False, font_color=None, bg_color=None, center_align=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            if font_color:
                run.font.color.rgb = font_color
        if center_align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bg_color:
        cell._element.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        )

def gen_tables(doc, num):
    paragraph = doc.add_paragraph()
    paragraph.add_run(f'2.{num} CVENUMBER')

    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'

    cell = table.cell(0, 0)
    cell_merge = cell.merge(table.cell(0, 3))
    cell.text = f'취약점 발생 정보'
    set_cell_style(cell, '맑은 고딕', 11, bold=True, font_color=RGBColor(255, 255, 255), bg_color="536061", center_align=True)

    table.cell(1, 0).text = f'URL{num}'
    table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_style(table.cell(1, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)

    table.cell(2, 0).text = f'Target-Key{num}'
    table.cell(2, 2).text = f'Target-Value{num}'
    set_cell_style(table.cell(2, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)
    set_cell_style(table.cell(2, 2), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)
    set_cell_style(table.cell(2, 1), '맑은 고딕', 10, bold=False, center_align=True)
    set_cell_style(table.cell(2, 3), '맑은 고딕', 10, bold=False, center_align=True)

    table.cell(3, 0).text = f'CVSS Score{num}'
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_style(table.cell(3, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)

    table.cell(4, 0).text = f'Payload{num}'
    table.cell(4, 1).merge(table.cell(4, 3))
    set_cell_style(table.cell(4, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)

    cell = table.cell(5, 0)
    cell_merge = cell.merge(table.cell(5, 3))
    cell.text = f'취약점 설명'
    set_cell_style(cell, '맑은 고딕', 11, bold=True, font_color=RGBColor(255, 255, 255), bg_color="536061", center_align=True)

    table.cell(6, 0).text = f'취약점{num}'
    table.cell(6, 1).merge(table.cell(6, 3))
    set_cell_style(table.cell(6, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)

    table.cell(7, 0).text = f'설명{num}'
    table.cell(7, 1).merge(table.cell(7, 3))
    set_cell_style(table.cell(7, 0), '맑은 고딕', 10, bold=True, font_color=RGBColor(0, 0, 0), bg_color="A7B3B5", center_align=True)

    cell = table.cell(8, 0)
    cell_merge = cell.merge(table.cell(8, 3))
    cell.text = f'대응 방안'
    set_cell_style(cell, '맑은 고딕', 11, bold=True, font_color=RGBColor(255, 255, 255), bg_color="536061", center_align=True)

    cell = table.cell(9, 0)
    cell_merge = cell.merge(table.cell(9, 3))
    cell.text = ''
    set_cell_style(cell, '맑은 고딕', 10, bold=False)

def gen_closing_page(doc):
    paragraph = doc.add_paragraph()
    
    run1 = paragraph.add_run("SHINNAM")
    run1.font.size = Pt(20)
    run1.font.name = '레시피코리아 Medium'
    run1.font.color.rgb = RGBColor(83, 96, 97)

    run2 = paragraph.add_run("\nContact. Shinnam@security.com")
    run2.font.size = Pt(10)
    run2.font.name = '레시피코리아 Medium'
    run2.font.color.rgb = RGBColor(83, 96, 97)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def update_toc(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)

            for run in paragraph.runs:
                run.font.name = '배달의민족 도현'
                run.font.size = Pt(10)