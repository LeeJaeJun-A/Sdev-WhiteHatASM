# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from datetime import datetime
import random
import os
import copy
from backend.database.mongodb import fs  # MongoDB GridFS
from backend.core.report_design import gen_tables, gen_closing_page, update_toc
import hashlib

doc = Document("/app/backend/core/reportv3.docx")

explain={
    "sqli":[
        {
            "vuln":"SQL Injeciton",
            "explain":"SQL 인젝션(SQL Injection)은 공격자가 웹 애플리케이션의 취약점을 이용해 악의적인 SQL 코드를 주입함으로써 데이터베이스를 비정상적으로 조작할 수 있는 보안 취약점이다. 이를 통해 공격자는 데이터베이스의 데이터를 조회, 수정, 삭제하거나, 관리자 권한을 탈취하는 등의 공격을 수행할 수 있다. SQL 인젝션은 주로 사용자 입력이 제대로 검증되지 않고 직접 SQL 쿼리에 포함될 때 발생한다. 예를 들어, 로그인 폼에 입력된 값이 그대로 SQL 쿼리에 사용되면, 공격자가 쿼리 구조를 변경하여 의도하지 않은 데이터베이스 명령을 실행할 수 있다.",
            "altplan":"소스코드에 SQL 쿼리를 입력 값으로 받는 함수나 코드를 사용할 경우, 임의의 SQL 쿼리 입력에 대한 검증 로직을 구현하여 서버에 검증되지 않는 SQL 쿼리 요청 시 에러 페이지가 아닌 정상 페이지가 반환되도록 필터링 처리하고 웹 방화벽에 SQL 인젝션 관련 룰셋을 적용하여 SQL 인젝션 공격을 차단한다.\n>> SQL 쿼리에 사용되는 문자열의 유효성을 검증하는 로직 구현\n>> Dynamic SQL 구문 사용을 지양하며 파라미터에 문자열 검사 필수적용\n>> 시스템에서 제공하는 에러 메시지 및 DBMS에서 제공하는 에러 코드가 노출되지 않도록 \n>> 아래와 같은 특수문자를 사용자 입력 값으로 지정 금지(아래 문자들은 해당 데이터베이스에 따라 달라질 수 있음)\n"
        }
    ],
    "xss":[
        {
            "vuln":"XSS(Cross-Site Scripting)",
            "explain":"XSS(Cross-Site Scripting)는 공격자가 웹 페이지에 악의적인 스크립트를 삽입하여 사용자의 브라우저에서 실행되도록 하는 보안 취약점이다. 이를 통해 공격자는 사용자의 세션 쿠키, 로그인 정보, 개인 데이터를 탈취하거나, 웹 페이지를 변조하는 등의 공격을 수행할 수 있다. XSS는 주로 사용자 입력을 제대로 검증하거나 이스케이프하지 않고 HTML이나 JavaScript 코드에 포함시킬 때 발생한다. 예를 들어, 댓글 폼에 입력된 스크립트가 그대로 출력되면, 해당 스크립트가 다른 사용자의 브라우저에서 실행될 수 있다.",
            "altplan":"웹 사이트의 게시판, 1:1 문의, URL 등에서 사용자 입력 값에 대해 검증 로직을 추가하거나 입력되더라도 실행되지 않게 하고, 부득이하게 웹페이지에서 HTML을 사용하는 경우 HTML 코드 중 필요한 코드에 대해서만 입력되게 설정한다.\n\n>> 웹 사이트에 사용자 입력 값이 저장되는 페이지는 공격자가 웹 브라우저를 통해 실행되는 스크립트 언어(HTML, Javascript, VBScript 등)를 사용하여 공격하므로 해당되는 태그 사용을 사전에 제한하고, 사용자 입력 값에 대한 필터링 작업이 필요함\n>> 게시물의 본문뿐만 아니라 제목, 댓글, 검색어 입력 창, 그 외 사용자 측에서 넘어오는 값을 신뢰하는 모든 form과 파라미터 값에 대해서 필터링을 수행함\n>> 입력 값에 대한 필터링 로직 구현 시 공백 문자를 제거하는 trim, replace 함수를 사용하여 반드시 서버 측에서 구현되어야 함\n>> URLDecoder 클래스에 존재하는 decode 메소드를 통해 URL 인코딩이 적용된 사용자 입력 값을 디코딩함으로써 우회 공격 차단\n>> 웹 방화벽에 모든 사용자 입력 폼(회원정보 변경, 게시판, 댓글, 자료실, 검색, URL 등)을 대상으로 특수문자, 특수 구문 필터링하도록 룰셋 적용\n\n※ 필터링 조치 대상 입력 값\n• 스크립트 정의어 : <script>, <object>, <applet>, <embed>, <form>, <iframe> 등\n• 특수문자 : <, >, “, ‘, &, %, %00(null) 등"
        }
    ],
    "csrf":[
        {
            "vuln":"CSRF(Cross-Site Request Forgery)",
            "explain":"CSRF(Cross Site Request Forgery): 사용자가 자신의 의지와는 무관하게 공격자가 의도한 행위(수정, 삭제, 등록 등)를 특정 웹 사이트에 요청하게 하는 공격 유형으로 사용자의 신뢰(인증) 정보 내에서 사용자의 요청(Request)을 변조함으로써 해당 사용자의 권한으로 악의적인 공격을 수행할 수 있다.",
            "altplan":"- 웹 사이트에 사용자 입력 값이 저장되는 페이지는 요청이 일회성이 될 수 있도록 설계\n- 사용 중인 프레임워크에 기본적으로 제공되는 CSRF 보호 기능 사용\n- 사용자가 정상적인 프로세스를 통해 요청하였는지 HTTP 헤더의 Referer 검증 로직 구현\n- 정상적인 요청(Request)과 비정상적인 요청(Request)를 구분할 수 있도록 Hidden Form을 사용하여 임의의 암호화된 토큰(세션 ID, Timestamp, nonce 등)을 추가하고 이 토큰을 검증하도록 설계\n- HTML이나 자바스크립트에 해당되는 태그 사용을 사전에 제한하고, 서버 단에서 사용자 입력 값에 대한 필터링 구현\n- HTML Editor 사용으로 인한 상기사항 조치 불가 시, 서버 사이드/서블릿/DAO\n(Data Access Object) 영역에서 조치하도록 설계\n- XSS 조치 방안 참조"
        }
    ],
    "xxe":[
        {
            "vuln":"XXE Injection",
            "explain":"XML 파싱 과정에서 외부 개체(External Entity)를 처리하는 기능을 악용하여 시스템 파일을 읽거나 서버의 비밀 정보를 유출시키는 공격 유형이 있다. XML 파서가 사용자 입력을 통해 전달받은 XML 데이터를 처리할 때 외부 개체를 참조하도록 허용되면, 공격자는 이를 악용하여 서버에서 로컬 파일을 읽거나 네트워크 요청을 수행하도록 유도할 수 있다. 이로 인해 서버 내부의 민감한 정보(예: /etc/passwd 파일 등)가 공격자에게 노출될 수 있으며, 경우에 따라 원격 서버로 데이터를 전송하거나, 서비스 거부(DoS) 공격을 유발할 수도 있다.",
            "altplan":"- XML 파서에서 외부 개체(Entity) 및 DTD(문서 유형 정의) 처리를 비활성화하도록 설정\n- 최신 보안 패치를 적용하여 XML 파서 및 관련 라이브러리 업데이트\n- 사용자 입력 값에 대한 철저한 검증 및 필터링을 통해 예상치 못한 외부 개체 참조 제거\n- 보안 분석 도구 및 웹 애플리케이션 방화벽(WAF)을 사용하여 XXE 취약점 자동 탐지 및 방어\n- 시스템 로그를 통해 XML 파서의 비정상적인 활동을 모니터링하고 신속히 대응"
        }
    ]
}

def apply_font_style(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)

def duplicate_section(doc, template_idx, vuln):
    template_paragraphs = doc.paragraphs[template_idx:template_idx+10]
    for para in template_paragraphs:
        new_para = doc.add_paragraph()
        new_para._element = copy.deepcopy(para._element)
        for run in new_para.runs:
            run.text = run.text.replace("CVENUMBER", vuln['cve'])
            run.text = run.text.replace("취약점 설명", vuln['description'])
            run.text = run.text.replace("대응 방안", vuln['remediation'])
            apply_font_style(run, '배달의민족 도현', 15)

def update_vulnerability_info(doc, table_index, start_idx, vulnerabilities, vuln, is_last=False):
    vulnerability_table = doc.tables[table_index]
    if not is_last:
        update_toc(doc, f"2.{table_index} CVELIST", f"2.{table_index} {vulnerabilities[0]['cve']}\n\t2.{table_index + 1} CVELIST")
    else:
        update_toc(doc, f"2.{table_index} CVELIST", f"2.{table_index} {vulnerabilities[0]['cve']}")

    for i, section in enumerate(vulnerabilities):
        row_idx = start_idx + i * 4
        
        if len(vulnerability_table.rows) <= row_idx + 5:
            for _ in range((row_idx + 6) - len(vulnerability_table.rows)):
                vulnerability_table.add_row()

        for paragraph in doc.paragraphs:
            if 'CVENUMBER' in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run('2.' + str(section['number']) + ' ' + section['cve'])
                apply_font_style(run, '배달의민족 도현', 15)
                break

        vulnerability_table.rows[row_idx].cells[1].text = section['url']

        vulnerability_table.rows[row_idx+1].cells[1].text = ', '.join(section['target_keys'])

        vulnerability_table.rows[row_idx+1].cells[3].text = ', '.join(section['target_values'])

        vulnerability_table.rows[row_idx+2].cells[1].text = section['cvss_score']

        vulnerability_table.rows[row_idx+3].cells[1].text = section['payload']

        description = explain[vuln][0]["vuln"]
        vulnerability_table.rows[row_idx+5].cells[1].text += description

        description = explain[vuln][0]["explain"]
        vulnerability_table.rows[row_idx+6].cells[1].text += description

        description = explain[vuln][0]["altplan"]
        vulnerability_table.rows[row_idx+8].cells[0].text += description

def find_template_section(doc, search_text):
    for i, paragraph in enumerate(doc.paragraphs):
        if search_text in paragraph.text:
            return i
    return -1

def gender_report(company_name, target_url, date, structure):
    # 1.1 수행 목적 업데이트
    for paragraph in doc.paragraphs:
        if '본 모의해킹 진단은' in paragraph.text:
            paragraph.text = paragraph.text.replace("company_name", f"'{company_name}'")
            break

    # 1.2 수행 대상 업데이트
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == '모의해킹':
                row.cells[1].text = target_url  
                row.cells[2].text = date 
                break 

    # 1.3 대상 구조 업데이트
    for paragraph in doc.paragraphs:
        if '대상구조적는곳' in paragraph.text:
            paragraph.text = paragraph.text.replace("대상구조적는곳", f"'{structure}'")
            break

# 문서 저장 및 MongoDB에 업로드
def save_report():
    gen_closing_page(doc)
    current_time = datetime.now().strftime("%Y%m%d%H%M%S%f")
    random_number = hashlib.sha256(current_time.encode()).hexdigest()[:8]
    os.makedirs("/app/backend/report", exist_ok=True)
    output_path = f"/app/backend/report/updated_{random_number}.docx"
    doc.save(output_path)

    with open(output_path, "rb") as file:
        #name=f"updated_{random_number}.docx"
        file_id = fs.put(file, filename=f"updated_{random_number}.docx")
        print(f"File uploaded to MongoDB with file_id: {file_id}")
        return file_id

def vuln_all(vulnerabilities):
    for i in range(len(vulnerabilities)):
        vuln = vulnerabilities[i]['vuln']
        is_last = (i == len(vulnerabilities) - 1)
        gen_tables(doc, vulnerabilities[i]['number'])
        update_vulnerability_info(doc, vulnerabilities[i]['number'], 1, [vulnerabilities[i]], vuln, is_last)
        doc.add_page_break()